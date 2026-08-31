import logging
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from typing import Annotated
from uuid import uuid4
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from fastapi import APIRouter, File, HTTPException, Query, Request, UploadFile, status
from lxml.etree import XMLSyntaxError  # type: ignore[import-untyped]
from starlette.concurrency import run_in_threadpool

from app.agents.messages import human_message
from app.agents.service import AnalysisDeskResult
from app.dependencies import ApplicationServices
from app.errors import AnalysisUnavailableError
from app.file_validation import content_matches_type
from app.models import (
    Asset,
    AssetUpload,
    Case,
    CaseSummary,
    Finding,
    ProductionMember,
    StoredAsset,
)
from app.models.requests import (
    ALLOWED_ANALYSIS_FILE_CONTENT_TYPES,
    ALLOWED_ASSET_CONTENT_TYPE,
    DOCX_CONTENT_TYPE,
    MAX_ANALYSIS_FILE_BYTES,
    MAX_ASSET_BYTES,
    MAX_EXTRACTED_DOCUMENT_CHARS,
    CreateCaseRequest,
    CreateThreadMessageRequest,
    UpdateFindingRequest,
)
from app.repositories import CaseRepositoryNotFound, FindingNotFound, ProductionRepositoryNotFound

router = APIRouter(prefix="/api/cases", tags=["cases"])
logger = logging.getLogger(__name__)
MAX_DOCX_ARCHIVE_FILES = 2_000
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024


def _services(request: Request) -> ApplicationServices:
    return request.app.state.services  # type: ignore[no-any-return]


async def _analyze_desk(
    services: ApplicationServices,
    case_id: str,
    script_text: str,
    ignored_keywords: list[str],
    roster: list[ProductionMember],
) -> AnalysisDeskResult:
    return await services.agent_service.analyze_desk(
        case_id, script_text, ignored_keywords, roster
    )


async def _analyze_file_desk(
    services: ApplicationServices,
    case_id: str,
    filename: str,
    content_type: str,
    content: bytes,
    ignored_keywords: list[str],
    roster: list[ProductionMember],
) -> AnalysisDeskResult:
    return await services.agent_service.analyze_file_desk(
        case_id, filename, content_type, content, ignored_keywords, roster
    )


async def _delete_asset_after_increment_failure(
    services: ApplicationServices, asset: StoredAsset
) -> None:
    try:
        await run_in_threadpool(services.asset_repository.delete, asset)
    except Exception:
        logger.warning("Asset cleanup failed after a case update error.")


def _extract_docx_text(content: bytes) -> str:
    try:
        with ZipFile(BytesIO(content)) as archive:
            entries = archive.infolist()
            if (
                len(entries) > MAX_DOCX_ARCHIVE_FILES
                or sum(entry.file_size for entry in entries) > MAX_DOCX_UNCOMPRESSED_BYTES
                or any(entry.flag_bits & 1 for entry in entries)
            ):
                raise HTTPException(
                    status_code=422,
                    detail="The DOCX archive is too large or encrypted",
                )
        document = Document(BytesIO(content))
    except (BadZipFile, KeyError, PackageNotFoundError, ValueError, XMLSyntaxError) as error:
        raise HTTPException(status_code=422, detail="The DOCX file could not be read") from error
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs]
    blocks.extend(
        cell.text.strip()
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    text = "\n".join(block for block in blocks if block)
    if not text:
        raise HTTPException(status_code=422, detail="The DOCX file does not contain readable text")
    if len(text) > MAX_EXTRACTED_DOCUMENT_CHARS:
        raise HTTPException(
            status_code=422,
            detail="The DOCX file contains too much text to analyze",
        )
    return text


@router.post("", response_model=Case, status_code=status.HTTP_201_CREATED)
async def create_case(payload: CreateCaseRequest, request: Request) -> Case:
    case_id = str(uuid4())
    services = _services(request)
    ignored_keywords: list[str] = []
    roster: list[ProductionMember] = []
    if payload.production_id is not None:
        try:
            production = await run_in_threadpool(
                services.production_repository.get, payload.production_id
            )
        except ProductionRepositoryNotFound as error:
            raise HTTPException(status_code=404, detail="Production not found") from error
        ignored_keywords = production.ignore_keywords
        roster = production.roster
    try:
        desk = await _analyze_desk(
            services, case_id, payload.script_text, ignored_keywords, roster
        )
    except AnalysisUnavailableError as error:
        logger.warning("Case analysis failed during %s.", error.operation)
        raise HTTPException(
            status_code=503,
            detail="RightsRadar analysis is temporarily unavailable. Please try again.",
        ) from error
    case = Case(
        id=case_id,
        script_text=payload.script_text,
        created_at=datetime.now(UTC),
        findings=desk.findings,
        production_id=payload.production_id,
        title=payload.title,
        thread=desk.thread,
        tool_calls=desk.tool_calls,
    )
    return await run_in_threadpool(services.case_repository.create, case)


@router.post(
    "/from-file/{production_id}",
    response_model=Case,
    status_code=status.HTTP_201_CREATED,
)
async def create_case_from_file(
    production_id: str,
    file: Annotated[UploadFile, File(json_schema_extra={"format": "binary"})],
    request: Request,
) -> Case:
    services = _services(request)
    try:
        production = await run_in_threadpool(
            services.production_repository.get, production_id
        )
    except ProductionRepositoryNotFound as error:
        raise HTTPException(status_code=404, detail="Production not found") from error

    content_type = file.content_type or ""
    if content_type not in ALLOWED_ANALYSIS_FILE_CONTENT_TYPES:
        raise HTTPException(
            status_code=422,
            detail="Analyze a PDF, DOCX, PNG, JPEG, or WebP file",
        )
    content = await file.read(MAX_ANALYSIS_FILE_BYTES + 1)
    if len(content) > MAX_ANALYSIS_FILE_BYTES:
        raise HTTPException(status_code=422, detail="Analysis files must not exceed 10 MiB")
    if not content_matches_type(content, content_type):
        raise HTTPException(
            status_code=422,
            detail="Analysis file content does not match its declared type",
        )

    filename = Path(file.filename or "production-file").name
    case_id = str(uuid4())
    try:
        if content_type == DOCX_CONTENT_TYPE:
            source_text = await run_in_threadpool(_extract_docx_text, content)
            desk = await _analyze_desk(
                services, case_id, source_text, production.ignore_keywords, production.roster
            )
            case_text = source_text[:20_000]
        else:
            desk = await _analyze_file_desk(
                services,
                case_id,
                filename,
                content_type,
                content,
                production.ignore_keywords,
                production.roster,
            )
            case_text = f"Uploaded {content_type} production file: {filename}"
    except AnalysisUnavailableError as error:
        logger.warning("File analysis failed during %s.", error.operation)
        raise HTTPException(
            status_code=503,
            detail="RightsRadar file analysis is temporarily unavailable. Please try again.",
        ) from error

    case = Case(
        id=case_id,
        script_text=case_text,
        created_at=datetime.now(UTC),
        findings=desk.findings,
        production_id=production_id,
        title=filename,
        thread=desk.thread,
        tool_calls=desk.tool_calls,
    )
    await run_in_threadpool(services.case_repository.create, case)
    asset: StoredAsset | None = None
    try:
        asset = await run_in_threadpool(
            services.asset_repository.store,
            case_id,
            AssetUpload(
                filename=filename,
                content_type=content_type,
                content=content,
            ),
        )
        await run_in_threadpool(services.case_repository.increment_asset_count, case_id)
    except Exception:
        if asset is not None:
            await _delete_asset_after_increment_failure(services, asset)
        try:
            await run_in_threadpool(services.case_repository.delete, case_id)
        except Exception:
            logger.warning("Case cleanup failed after analyzed file persistence failed.")
        raise
    return await run_in_threadpool(services.case_repository.get, case_id)


@router.get("", response_model=list[CaseSummary])
def list_cases(request: Request, limit: int = Query(default=10, ge=1, le=50)) -> list[CaseSummary]:
    return _services(request).case_repository.list_recent(limit)


@router.get("/{case_id}", response_model=Case)
def get_case(case_id: str, request: Request) -> Case:
    try:
        return _services(request).case_repository.get(case_id)
    except CaseRepositoryNotFound as error:
        raise HTTPException(status_code=404, detail="Case not found") from error


@router.delete("/{case_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_case(case_id: str, request: Request) -> None:
    services = _services(request)
    try:
        services.case_repository.get(case_id)
    except CaseRepositoryNotFound as error:
        raise HTTPException(status_code=404, detail="Case not found") from error
    for asset in services.asset_repository.list_for_case(case_id):
        try:
            services.asset_repository.delete(asset)
        except Exception:
            logger.warning("Asset cleanup failed while deleting case %s.", case_id)
    services.case_repository.delete(case_id)


@router.post("/{case_id}/assets", response_model=Asset, status_code=status.HTTP_201_CREATED)
async def upload_asset(
    case_id: str,
    file: Annotated[UploadFile, File(json_schema_extra={"format": "binary"})],
    request: Request,
) -> Asset:
    services = _services(request)
    try:
        await run_in_threadpool(services.case_repository.get, case_id)
    except CaseRepositoryNotFound as error:
        raise HTTPException(status_code=404, detail="Case not found") from error

    if file.content_type != ALLOWED_ASSET_CONTENT_TYPE:
        raise HTTPException(status_code=422, detail="Only text/plain assets are supported")
    content = await file.read(MAX_ASSET_BYTES + 1)
    if len(content) > MAX_ASSET_BYTES:
        raise HTTPException(status_code=422, detail="Asset must not exceed 256 KiB")
    if b"\x00" in content:
        raise HTTPException(status_code=422, detail="Asset must contain valid UTF-8 text")
    try:
        content.decode("utf-8")
    except UnicodeDecodeError as error:
        raise HTTPException(
            status_code=422, detail="Asset must contain valid UTF-8 text"
        ) from error

    asset = await run_in_threadpool(
        services.asset_repository.store,
        case_id,
        AssetUpload(
            filename=file.filename or "asset.txt",
            content_type=ALLOWED_ASSET_CONTENT_TYPE,
            content=content,
        ),
    )
    try:
        await run_in_threadpool(services.case_repository.increment_asset_count, case_id)
    except CaseRepositoryNotFound as error:
        await _delete_asset_after_increment_failure(services, asset)
        raise HTTPException(status_code=404, detail="Case not found") from error
    except Exception:
        await _delete_asset_after_increment_failure(services, asset)
        raise
    return asset


@router.get("/{case_id}/assets", response_model=list[Asset])
def list_assets(case_id: str, request: Request) -> list[Asset]:
    services = _services(request)
    try:
        services.case_repository.get(case_id)
    except CaseRepositoryNotFound as error:
        raise HTTPException(status_code=404, detail="Case not found") from error
    return [
        Asset.model_validate(asset) for asset in services.asset_repository.list_for_case(case_id)
    ]


@router.post("/{case_id}/thread", response_model=Case, status_code=status.HTTP_201_CREATED)
def post_thread_message(
    case_id: str, payload: CreateThreadMessageRequest, request: Request
) -> Case:
    services = _services(request)
    try:
        case = services.case_repository.get(case_id)
    except CaseRepositoryNotFound as error:
        raise HTTPException(status_code=404, detail="Case not found") from error
    if payload.finding_id and not any(
        finding.id == payload.finding_id for finding in case.findings
    ):
        raise HTTPException(status_code=404, detail="Finding not found")
    member_id = payload.member_id
    if case.production_id:
        try:
            production = services.production_repository.get(case.production_id)
        except ProductionRepositoryNotFound as error:
            raise HTTPException(status_code=404, detail="Production not found") from error
        if not any(member.id == member_id for member in production.roster):
            raise HTTPException(status_code=422, detail="Member is not on this production roster")
    try:
        return services.case_repository.add_thread_message(
            case_id,
            human_message(
                case_id,
                member_id,
                payload.body,
                finding_id=payload.finding_id,
            ),
        )
    except CaseRepositoryNotFound as error:
        raise HTTPException(status_code=404, detail="Case not found") from error


@router.patch("/{case_id}/findings/{finding_id}", response_model=Finding)
def update_finding(
    case_id: str, finding_id: str, payload: UpdateFindingRequest, request: Request
) -> Finding:
    services = _services(request)
    try:
        finding = services.case_repository.update_finding_status(
            case_id, finding_id, payload.reviewer_status
        )
    except CaseRepositoryNotFound as error:
        raise HTTPException(status_code=404, detail="Case not found") from error
    except FindingNotFound as error:
        raise HTTPException(status_code=404, detail="Finding not found") from error
    if payload.actor_member_id:
        case = services.case_repository.get(case_id)
        actor_name = payload.actor_member_id
        actor_role = "reviewer"
        if case.production_id:
            try:
                production = services.production_repository.get(case.production_id)
            except ProductionRepositoryNotFound:
                production = None
            if production is not None:
                actor = next(
                    (
                        member
                        for member in production.roster
                        if member.id == payload.actor_member_id
                    ),
                    None,
                )
                if actor is not None:
                    actor_name = actor.name
                    actor_role = actor.role.value
        action = (
            "escalated"
            if payload.reviewer_status.value == "escalated"
            else (
                "dismissed"
                if payload.reviewer_status.value == "dismissed"
                else payload.reviewer_status.value
            )
        )
        services.case_repository.add_thread_message(
            case_id,
            human_message(
                case_id,
                payload.actor_member_id,
                (
                    f"{actor_name} ({actor_role}) {action} {finding.detected_item} "
                    "in the case desk thread."
                ),
                finding_id=finding_id,
            ),
        )
    return finding
